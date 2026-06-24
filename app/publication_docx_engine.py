"""
Publication DOCX Engine — Génère publication.docx depuis publication.md.

Entrée  : sortie/<project_name>/publication/publication.md
Sortie  : sortie/<project_name>/publication/publication.docx

Moteur modulaire conçu pour supporter à terme :
  - couverture graphique
  - PDF professionnel
  - guide du participant / formateur
  - cahier d'exercices
  - livre / livret
"""

from __future__ import annotations

import re
from datetime import date, datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.document_language import get_document_language
from app.logger import log_event
from app.paths import SORTIE_DIR
from app.project_state import load_project_state, save_project_state
from app.publication_metadata import get_publication_metadata
from app.publication_theme import get_publication_theme


# ─────────────────────────────────────────────────────────────────────────────
# Labels par langue
# ─────────────────────────────────────────────────────────────────────────────

_LABELS: dict[str, dict[str, str]] = {
    "en": {
        "cover":          "Cover",
        "title_page":     "Title Page",
        "toc":            "Table of Contents",
        "title_label":    "Title",
        "project_label":  "Project",
        "date_label":     "Date",
    },
    "fr": {
        "cover":          "Couverture",
        "title_page":     "Page de titre",
        "toc":            "Table des matières",
        "title_label":    "Titre",
        "project_label":  "Projet",
        "date_label":     "Date",
    },
}


# ─────────────────────────────────────────────────────────────────────────────
# Helpers DOCX bas niveau
# ─────────────────────────────────────────────────────────────────────────────

def _add_page_break(doc: Document) -> None:
    """Insère un saut de page DOCX natif."""
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _style_title(doc: Document, font_size: float = 28) -> None:
    """Configure le style Title : centré, grand, gras."""
    style = doc.styles["Title"]
    font = style.font
    font.size = Pt(font_size)
    font.bold = True
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style.paragraph_format.space_before = Pt(24)
    style.paragraph_format.space_after = Pt(12)


def _style_heading1(doc: Document, font_size: float = 18) -> None:
    """Configure Heading 1 : gras, couleur primaire, espace."""
    style = doc.styles["Heading 1"]
    font = style.font
    font.size = Pt(font_size)
    font.bold = True
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    style.paragraph_format.space_before = Pt(18)
    style.paragraph_format.space_after = Pt(6)


def _style_heading2(doc: Document, font_size: float = 14) -> None:
    """Configure Heading 2 : semi-gras, couleur secondaire."""
    style = doc.styles["Heading 2"]
    font = style.font
    font.size = Pt(font_size)
    font.bold = True
    font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    style.paragraph_format.space_before = Pt(12)
    style.paragraph_format.space_after = Pt(4)


def _style_normal(doc: Document, font_size: float = 11, line_spacing: float = 1.15) -> None:
    """Configure Normal : lisible, interligné confortable."""
    style = doc.styles["Normal"]
    font = style.font
    font.size = Pt(font_size)
    font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = Pt(font_size * line_spacing * 1.2)


def _apply_styles(doc: Document, theme: dict | None = None) -> None:
    t = theme or {}
    _style_title(doc,    font_size=t.get("title_font_size",    28))
    _style_heading1(doc, font_size=t.get("heading1_font_size", 18))
    _style_heading2(doc, font_size=t.get("heading2_font_size", 14))
    _style_normal(doc,
                  font_size=t.get("body_font_size",    11),
                  line_spacing=t.get("body_line_spacing", 1.15))


def _apply_margins(doc: Document, theme: dict) -> None:
    """Applique les marges de page depuis le thème."""
    section = doc.sections[0]
    section.top_margin    = Pt(theme.get("top_margin",    72))
    section.bottom_margin = Pt(theme.get("bottom_margin", 72))
    section.left_margin   = Pt(theme.get("left_margin",   72))
    section.right_margin  = Pt(theme.get("right_margin",  72))


# ─────────────────────────────────────────────────────────────────────────────
# Inline Markdown dans les runs
# ─────────────────────────────────────────────────────────────────────────────

_INLINE_RE = re.compile(r"(\*\*(.+?)\*\*|_(.+?)_|\*(.+?)\*|`(.+?)`)")

# Supprime les commentaires HTML Markdown (<!-- ... -->) avant le rendu DOCX.
# Ces commentaires sont utilisés par publication_builder comme marqueurs de
# sections vides et ne doivent jamais apparaître dans les livrables clients.
_HTML_COMMENT_RE = re.compile(r"<!--.*?-->", re.DOTALL)


def _add_inline_runs(para, text: str) -> None:
    """
    Ajoute des runs dans `para` en interprétant les marqueurs inline :
      **bold**  _italic_  *italic*  `code`
    """
    last = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > last:
            para.add_run(text[last:m.start()])
        raw = m.group(0)
        if raw.startswith("**"):
            run = para.add_run(m.group(2))
            run.bold = True
        elif raw.startswith("`"):
            run = para.add_run(m.group(5))
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        else:
            run = para.add_run(m.group(3) or m.group(4))
            run.italic = True
        last = m.end()
    if last < len(text):
        para.add_run(text[last:])


# ─────────────────────────────────────────────────────────────────────────────
# Pages structurelles
# ─────────────────────────────────────────────────────────────────────────────

def _add_cover_page(
    doc: Document,
    title: str,
    project_name: str,
    labels: dict[str, str],
    metadata: dict | None = None,
    cover_png_path: Path | None = None,
) -> None:
    """Ajoute la page de couverture (première page, pas de saut avant).

    Si cover_png_path est fourni et existe, insère l'image en couverture
    à la place de la couverture textuelle.
    """
    # ── Couverture image (cover.png généré par cover_builder) ─────────────
    if cover_png_path and cover_png_path.exists():
        try:
            section = doc.sections[0]
            # Largeur de la zone de contenu (hors marges)
            content_w = section.page_width - section.left_margin - section.right_margin
            img_para = doc.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_img = img_para.add_run()
            run_img.add_picture(str(cover_png_path), width=content_w)
            _add_page_break(doc)
            return
        except Exception as exc:
            print(f"[publication_docx_engine] Impossible d'insérer cover.png : {exc}")
            # Fallback vers la couverture textuelle ci-dessous

    meta     = metadata or {}
    pub_date = meta.get("publication_date") or date.today().isoformat()
    subtitle = meta.get("subtitle", "")
    author   = meta.get("author", "")
    org      = meta.get("organization", "")

    doc.add_paragraph(title, style="Title")

    if subtitle:
        sub_para = doc.add_paragraph(style="Normal")
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_para.paragraph_format.space_before = Pt(12)
        run_sub = sub_para.add_run(subtitle)
        run_sub.italic = True

    if author:
        auth_para = doc.add_paragraph(style="Normal")
        auth_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        auth_para.paragraph_format.space_before = Pt(8)
        auth_para.add_run(author)

    if org:
        org_para = doc.add_paragraph(style="Normal")
        org_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        org_para.add_run(org)

    meta_para = doc.add_paragraph(style="Normal")
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_para.paragraph_format.space_before = Pt(32)
    meta_para.add_run(f"{labels['date_label']} : {pub_date}")

    _add_page_break(doc)


def _add_title_page(
    doc: Document,
    title: str,
    project_name: str,
    labels: dict[str, str],
    metadata: dict | None = None,
) -> None:
    """Ajoute la page de titre (nouvelle page)."""
    meta     = metadata or {}
    pub_date = meta.get("publication_date") or date.today().isoformat()
    subtitle = meta.get("subtitle", "")
    author   = meta.get("author", "")
    org      = meta.get("organization", "")

    heading = doc.add_paragraph(style="Title")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_before = Pt(80)
    heading.add_run(title)

    if subtitle:
        sub_para = doc.add_paragraph(style="Normal")
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_para.paragraph_format.space_before = Pt(16)
        run_sub = sub_para.add_run(subtitle)
        run_sub.italic = True

    if author:
        auth_para = doc.add_paragraph(style="Normal")
        auth_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        auth_para.paragraph_format.space_before = Pt(24)
        auth_para.add_run(author)

    if org:
        org_para = doc.add_paragraph(style="Normal")
        org_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_org = org_para.add_run(org)
        run_org.italic = True

    date_para = doc.add_paragraph(style="Normal")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.paragraph_format.space_before = Pt(12)
    date_para.add_run(pub_date)

    _add_page_break(doc)


def _add_toc_page(
    doc: Document,
    toc_entries: list[str],
    labels: dict[str, str],
) -> None:
    """Ajoute la table des matières (nouvelle page)."""
    doc.add_heading(labels["toc"], level=1)

    for entry in toc_entries:
        para = doc.add_paragraph(style="Normal")
        para.paragraph_format.left_indent = Inches(0.3)
        _add_inline_runs(para, entry)

    _add_page_break(doc)


# ─────────────────────────────────────────────────────────────────────────────
# Parseur publication.md → entrées TOC + corps
# ─────────────────────────────────────────────────────────────────────────────

_STRUCTURAL_KEYS = {
    "cover", "couverture",
    "title page", "page de titre",
    "table of contents", "table des matières",
}

_HEADING_RE = re.compile(r"^(#{1,3})\s+(.+)$")
_HR_RE = re.compile(r"^---+$")
_BULLET_RE = re.compile(r"^(\s*)([*\-+]|\d+\.)\s+(.+)$")


def _extract_title_from_md(lines: list[str]) -> str:
    """
    Récupère le premier texte en gras ou le premier H2 après la couverture.
    Fallback : "Publication".
    """
    in_cover = False
    for line in lines:
        stripped = line.strip()
        m = _HEADING_RE.match(stripped)
        if m:
            key = m.group(2).strip().lower()
            if key in ("cover", "couverture"):
                in_cover = True
                continue
            if key in _STRUCTURAL_KEYS:
                in_cover = False
                continue
        if in_cover and stripped.startswith("**") and stripped.endswith("**"):
            return stripped[2:-2].strip()
    return "Publication"


def _parse_toc_entries(lines: list[str]) -> list[str]:
    """
    Extrait les entrées de la table des matières depuis publication.md.
    Section délimitée par la ligne `# Table of Contents` (ou FR équiv.)
    et la prochaine ligne `---`.
    """
    entries: list[str] = []
    in_toc = False

    for line in lines:
        stripped = line.strip()
        if _HR_RE.match(stripped):
            if in_toc:
                break
            continue
        m = _HEADING_RE.match(stripped)
        if m:
            key = m.group(2).strip().lower()
            if key in ("table of contents", "table des matières"):
                in_toc = True
                continue
            if in_toc:
                break
        if in_toc and stripped:
            entries.append(stripped)

    return entries


def _convert_body(doc: Document, lines: list[str]) -> None:
    """
    Convertit les lignes du corps de publication.md en éléments DOCX.
    Ignore les sections structurelles (cover, title page, TOC).
    Supporte H1, H2, H3, puces, listes numérotées, paragraphes normaux.
    """
    skip_structural = True
    structural_count = 0  # on ignore les 3 premières sections structurelles

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Séparateur horizontal → ignore
        if _HR_RE.match(stripped):
            i += 1
            continue

        # Titres
        m = _HEADING_RE.match(stripped)
        if m:
            level = len(m.group(1))
            title_text = m.group(2).strip()
            key = title_text.lower()

            if key in _STRUCTURAL_KEYS:
                structural_count += 1
                i += 1
                continue

            if level == 1:
                doc.add_heading(title_text, level=1)
            elif level == 2:
                doc.add_heading(title_text, level=2)
            else:
                doc.add_heading(title_text, level=3)
            i += 1
            continue

        # Ligne vide → espace entre paragraphes (pas de para vide)
        if not stripped:
            i += 1
            continue

        # Puces et listes numérotées
        bm = _BULLET_RE.match(line)
        if bm:
            content = bm.group(3).strip()
            para = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(para, content)
            i += 1
            continue

        # Paragraphe normal
        para = doc.add_paragraph(style="Normal")
        _add_inline_runs(para, stripped)
        i += 1


# ─────────────────────────────────────────────────────────────────────────────
# Point d'entrée public
# ─────────────────────────────────────────────────────────────────────────────

def generate_publication_docx(project_name: str) -> Path | None:
    """
    Génère publication.docx depuis publication.md.

    Entrée  : sortie/<project_name>/publication/publication.md
    Sortie  : sortie/<project_name>/publication/publication.docx

    Retourne le chemin du DOCX généré, ou None en cas d'erreur.
    """
    pub_dir  = SORTIE_DIR / project_name / "publication"
    md_path  = pub_dir / "publication.md"
    docx_path = pub_dir / "publication.docx"
    now = datetime.now().isoformat(timespec="seconds")

    print(f"[publication_docx_engine] Génération DOCX — projet : {project_name}")
    log_event({
        "step": "publication_docx_engine",
        "project": project_name,
        "action": "start",
    })

    # ── Sélection de la source : préférer la version sanitized ───────────────
    sanitized_path = pub_dir / "publication_sanitized.md"
    effective_path = sanitized_path if sanitized_path.exists() else md_path

    if not effective_path.exists():
        msg = f"publication.md introuvable : {md_path}"
        print(f"[publication_docx_engine] ERREUR : {msg}")
        log_event({
            "step": "publication_docx_engine",
            "project": project_name,
            "action": "error",
            "error": msg,
        })
        _save_docx_state(project_name, generated=False, error=msg, now=now)
        return None

    _source_label = "publication_sanitized.md" if sanitized_path.exists() else "publication.md"
    print(f"[publication_docx_engine] Publication utilisée : {effective_path}  [{_source_label}]")
    log_event({
        "step":   "publication_docx_engine",
        "project": project_name,
        "action": "source",
        "path":   str(effective_path),
        "source": _source_label,
    })

    # ── Lecture ───────────────────────────────────────────────────────────
    md_text = effective_path.read_text(encoding="utf-8")
    # Supprimer les éventuels commentaires HTML résiduels avant tout traitement
    md_text = _HTML_COMMENT_RE.sub("", md_text)
    lines = md_text.splitlines()

    # ── Langue documentaire ───────────────────────────────────────────────
    lang = get_document_language(project_name, fallback_text=md_text)
    labels = _LABELS.get(lang, _LABELS["en"])
    print(f"[publication_docx_engine] Langue documentaire : {lang}")

    # ── Récupération du mode de publication ───────────────────────────────
    state = load_project_state(project_name)
    publication_mode = state.get("publication_mode") or "BOOK"
    print(f"[publication_docx_engine] Mode de publication : {publication_mode}")

    # ── Chargement du thème ───────────────────────────────────────────────
    theme = get_publication_theme(publication_mode, lang)
    print(
        f"[publication_docx_engine] Thème chargé : {theme['mode']} | "
        f"cover_style={theme['cover_style']} | "
        f"body={theme['body_font_size']}pt | "
        f"h1={theme['heading1_font_size']}pt"
    )
    log_event({
        "step":    "publication_docx_engine",
        "project": project_name,
        "action":  "theme_loaded",
        "mode":    theme["mode"],
        "cover_style": theme["cover_style"],
        "body_font_size": theme["body_font_size"],
    })

    # ── Métadonnées de publication ────────────────────────────────────────
    pub_meta = get_publication_metadata(project_name)
    print(
        f"[publication_docx_engine] Métadonnées appliquées au DOCX — "
        f"titre={pub_meta['title']!r} | auteur={pub_meta['author']!r}"
    )
    log_event({
        "step":    "publication_docx_engine",
        "project": project_name,
        "action":  "metadata_applied",
        "title":   pub_meta["title"],
        "author":  pub_meta["author"],
    })

    # ── Extraction titre + TOC ────────────────────────────────────────────
    # Le titre des métadonnées prend la priorité sur celui extrait du MD
    title_from_meta = pub_meta.get("title", "").strip()
    title_from_md   = _extract_title_from_md(lines)
    title           = title_from_meta if title_from_meta else title_from_md
    toc_entries     = _parse_toc_entries(lines)

    print(f"[publication_docx_engine] Titre : « {title} »")
    print(f"[publication_docx_engine] TOC : {len(toc_entries)} entrée(s)")

    # ── Construction du document ──────────────────────────────────────────
    try:
        doc = Document()
        _apply_styles(doc, theme)
        _apply_margins(doc, theme)

        # Propriétés internes du document DOCX
        core          = doc.core_properties
        core.title    = title
        core.author   = pub_meta.get("author") or "TranscriptionAI / PublishForge"
        core.subject  = pub_meta.get("subtitle") or f"Publication — {project_name}"
        core.keywords = f"{project_name}, {publication_mode}, {lang}"
        core.category = publication_mode
        if pub_meta.get("organization"):
            core.company = pub_meta["organization"]

        # 1. Couverture (première page)
        if theme.get("include_cover", True):
            # cover_image.png (moteur image externe) a la priorité sur cover.png (cover_builder)
            _cover_image_png = (
                SORTIE_DIR / project_name / "publication" / "cover" / "cover_image.png"
            )
            _cover_png = (
                SORTIE_DIR / project_name / "publication" / "cover" / "cover.png"
            )
            _effective_cover = (
                _cover_image_png if _cover_image_png.exists() else _cover_png
            )
            _add_cover_page(doc, title, project_name, labels, pub_meta, _effective_cover)

        # 2. Page de titre (nouvelle page)
        if theme.get("include_title_page", True):
            _add_title_page(doc, title, project_name, labels, pub_meta)

        # 3. Table des matières (nouvelle page)
        if theme.get("include_toc", True):
            _add_toc_page(doc, toc_entries, labels)

        # 4. Corps du document
        _convert_body(doc, lines)

        # ── Écriture ──────────────────────────────────────────────────────
        pub_dir.mkdir(parents=True, exist_ok=True)
        doc.save(str(docx_path))

    except Exception as exc:  # noqa: BLE001
        msg = str(exc)
        print(f"[publication_docx_engine] ERREUR génération : {msg}")
        log_event({
            "step": "publication_docx_engine",
            "project": project_name,
            "action": "error",
            "error": msg,
        })
        _save_docx_state(project_name, generated=False, error=msg, now=now)
        return None

    print(f"[publication_docx_engine] DOCX généré : {docx_path}")
    log_event({
        "step":    "publication_docx_engine",
        "project": project_name,
        "action":  "generated",
        "path":    str(docx_path),
        "mode":    publication_mode,
        "lang":    lang,
    })

    _save_docx_state(
        project_name,
        generated=True,
        path=str(docx_path),
        now=now,
        theme=publication_mode,
    )
    return docx_path


# ─────────────────────────────────────────────────────────────────────────────
# Mise à jour du project_state
# ─────────────────────────────────────────────────────────────────────────────

def _save_docx_state(
    project_name: str,
    *,
    generated: bool,
    now: str,
    path: str | None = None,
    error: str | None = None,
    theme: str | None = None,
) -> None:
    state = load_project_state(project_name)
    state.setdefault("publication", {})

    entry: dict = {"generated": generated, "generated_at": now}
    if generated and path:
        entry["path"] = path
    if not generated and error:
        entry["error"] = error
    if theme:
        entry["theme"] = theme

    state["publication"]["docx_engine"] = entry
    save_project_state(project_name, state)

    print(f"[publication_docx_engine] project_state.json mis à jour.")
