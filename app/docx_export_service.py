from pathlib import Path
from datetime import datetime
import hashlib
import json
import re

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.paths import SORTIE_DIR
from app.project_state import load_project_state, save_project_state


def _file_hash(path: Path) -> str:
    hasher = hashlib.md5()
    with path.open("rb") as f:
        for block in iter(lambda: f.read(8192), b""):
            hasher.update(block)
    return hasher.hexdigest()


def _get_final_dir(project_name: str) -> Path:
    return SORTIE_DIR / project_name / "final"


def _get_md_path(project_name: str) -> Path:
    return _get_final_dir(project_name) / "document_final.md"


def _get_docx_path(project_name: str) -> Path:
    return _get_final_dir(project_name) / "document_final.docx"


def _should_regenerate(
    state: dict,
    md_path: Path,
    docx_path: Path,
    current_signature: str,
) -> bool:
    if not docx_path.exists():
        return True

    exports = state.get("exports", {})
    docx_state = exports.get("docx", {})

    if not docx_state.get("generated"):
        return True

    if docx_state.get("source_signature") != current_signature:
        return True

    return False


def _set_document_metadata(doc: Document, project_name: str) -> None:
    core_props = doc.core_properties
    core_props.title = f"TranscriptionAI - {project_name}"
    core_props.author = "TranscriptionAI"
    core_props.subject = "Document généré automatiquement"


def _add_code_block(doc: Document, text: str) -> None:
    for line in text.splitlines():
        para = doc.add_paragraph(line)
        for run in para.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        if not para.runs:
            run = para.add_run(line)
            run.font.name = "Courier New"
            run.font.size = Pt(10)


def _convert_markdown_to_docx(md_text: str, doc: Document) -> None:
    """
    Parse le Markdown ligne par ligne et construit le document Word.

    Éléments supportés :
      - Titres H1 / H2 / H3
      - Listes à puces (* ou -)
      - Blocs de citation (>)
      - Blocs de code (``` ... ```)
      - Séparateurs horizontaux (---)
      - Paragraphes normaux
    """
    lines = md_text.splitlines()
    i = 0

    while i < len(lines):
        line = lines[i]

        # Bloc de code (``` ... ```)
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            _add_code_block(doc, "\n".join(code_lines))
            i += 1
            continue

        # Séparateur horizontal
        stripped = line.strip()
        if re.match(r"^-{3,}$", stripped) or re.match(r"^\*{3,}$", stripped):
            doc.add_paragraph("─" * 50)
            i += 1
            continue

        # Titre H3
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue

        # Titre H2
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue

        # Titre H1
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue

        # Liste à puces (* ou -)
        if re.match(r"^[*\-] ", stripped):
            text = re.sub(r"^[*\-] ", "", stripped)
            doc.add_paragraph(text, style="List Bullet")
            i += 1
            continue

        # Bloc de citation (>)
        if stripped.startswith("> "):
            text = stripped[2:].strip()
            try:
                doc.add_paragraph(text, style="Quote")
            except KeyError:
                doc.add_paragraph(text)
            i += 1
            continue

        # Ligne vide → séparation de paragraphe (on la saute)
        if stripped == "":
            i += 1
            continue

        # Paragraphe normal
        doc.add_paragraph(stripped)
        i += 1


def export_docx(project_name: str) -> Path | None:
    """
    Génère document_final.docx à partir de document_final.md.

    Régénère uniquement si :
      - le DOCX n'existe pas
      - le MD est plus récent (signature changée)
      - la signature enregistrée ne correspond plus

    Retourne le chemin du DOCX, ou None si le MD source est absent.
    """
    md_path = _get_md_path(project_name)
    docx_path = _get_docx_path(project_name)

    if not md_path.exists():
        print(
            f"[docx_export] document_final.md introuvable pour le projet : {project_name}"
        )
        return None

    current_signature = _file_hash(md_path)
    state = load_project_state(project_name)

    if not _should_regenerate(state, md_path, docx_path, current_signature):
        print(f"[docx_export] DOCX déjà à jour : {docx_path}")
        return docx_path

    md_text = md_path.read_text(encoding="utf-8")

    doc = Document()
    _set_document_metadata(doc, project_name)
    _convert_markdown_to_docx(md_text, doc)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))

    updated_at = datetime.now().isoformat(timespec="seconds")

    if "exports" not in state:
        state["exports"] = {}

    state["exports"]["docx"] = {
        "generated": True,
        "path": str(docx_path),
        "updated_at": updated_at,
        "source_signature": current_signature,
    }

    save_project_state(project_name, state)

    print(f"[docx_export] DOCX généré : {docx_path}")

    return docx_path


# ===========================================================================
# Publication DOCX — export professionnel avec gabarits et thèmes
# ===========================================================================

_PAGE_SIZES_DOCX: dict[str, tuple] = {
    "letter":      (Inches(8.5),  Inches(11.0)),
    "a4":          (Inches(8.27), Inches(11.69)),
    "six_by_nine": (Inches(6.0),  Inches(9.0)),
    "digest":      (Inches(5.5),  Inches(8.5)),
}

_PAGE_MARGINS_DOCX: dict[str, float] = {
    "letter":      1.0,
    "a4":          1.0,
    "six_by_nine": 0.75,
    "digest":      0.75,
}

_FONT_CONFIGS_DOCX: dict[str, dict] = {
    "classic": {
        "heading": "Times New Roman",
        "body":    "Georgia",
        "h1_size": 22, "h2_size": 16, "h3_size": 13, "body_size": 11,
        "h1_bold": True, "h2_bold": True, "h3_bold": True,
    },
    "modern": {
        "heading": "Calibri",
        "body":    "Calibri",
        "h1_size": 22, "h2_size": 16, "h3_size": 13, "body_size": 11,
        "h1_bold": True, "h2_bold": True, "h3_bold": False,
    },
    "elegant": {
        "heading": "Times New Roman",
        "body":    "Calibri",
        "h1_size": 24, "h2_size": 18, "h3_size": 14, "body_size": 12,
        "h1_bold": True, "h2_bold": True, "h3_bold": False,
    },
    "readable": {
        "heading": "Arial",
        "body":    "Arial",
        "h1_size": 20, "h2_size": 15, "h3_size": 12, "body_size": 12,
        "h1_bold": True, "h2_bold": True, "h3_bold": False,
    },
}


def _hex_to_rgb(hex_color: str) -> tuple[int, int, int]:
    h = hex_color.lstrip("#")
    return (int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _apply_page_size_docx(doc: Document, page_size_key: str) -> None:
    width, height = _PAGE_SIZES_DOCX.get(page_size_key, _PAGE_SIZES_DOCX["letter"])
    margin = Inches(_PAGE_MARGINS_DOCX.get(page_size_key, 1.0))
    section = doc.sections[0]
    section.page_width = width
    section.page_height = height
    section.left_margin = margin
    section.right_margin = margin
    section.top_margin = margin
    section.bottom_margin = margin


def _add_page_number_field(paragraph) -> None:
    """Insère un champ PAGE (numéro de page Word) dans un paragraphe."""
    run = paragraph.add_run()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar_begin)

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    run._r.append(instrText)

    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar_end)


def _add_headers_footers_docx(
    doc: Document,
    settings: dict,
    font_cfg: dict,
) -> None:
    """Ajoute en-tête et pied de page si activés dans les paramètres."""
    section = doc.sections[0]
    gray = "#888888"

    if settings.get("include_headers", True):
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.clear()
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hp.add_run(settings.get("title", ""))
        run.font.name = font_cfg["heading"]
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(*_hex_to_rgb(gray))

    if settings.get("include_footers", True):
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        org = settings.get("organization", "")

        if settings.get("include_page_numbers", True):
            # Organisation — Page N
            prefix = f"{org}  —  Page " if org else "Page "
            run_pre = fp.add_run(prefix)
            run_pre.font.name = font_cfg["body"]
            run_pre.font.size = Pt(9)
            run_pre.font.color.rgb = RGBColor(*_hex_to_rgb(gray))
            _add_page_number_field(fp)
        else:
            run = fp.add_run(org)
            run.font.name = font_cfg["body"]
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(*_hex_to_rgb(gray))


def _add_cover_page_docx(doc: Document, settings: dict, font_cfg: dict) -> None:
    """Génère une page de couverture typographique professionnelle."""
    colors_theme = settings.get("theme_colors", {})
    primary = colors_theme.get("primary", "#1a1a1a")
    accent = colors_theme.get("accent", "#666666")
    gray = "#aaaaaa"

    # Espace vertical supérieur (~30% de page)
    for _ in range(8):
        doc.add_paragraph()

    # Titre principal
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(settings.get("title", ""))
    title_run.font.name = font_cfg["heading"]
    title_run.font.size = Pt(32)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(*_hex_to_rgb(primary))

    # Sous-titre
    subtitle = settings.get("subtitle", "")
    if subtitle:
        doc.add_paragraph()
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_para.add_run(subtitle)
        sub_run.font.name = font_cfg["heading"]
        sub_run.font.size = Pt(16)
        sub_run.italic = True
        sub_run.font.color.rgb = RGBColor(*_hex_to_rgb(accent))

    # Séparateur décoratif
    doc.add_paragraph()
    sep_para = doc.add_paragraph()
    sep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep_run = sep_para.add_run("― ✦ ―")
    sep_run.font.name = font_cfg["heading"]
    sep_run.font.size = Pt(14)
    sep_run.font.color.rgb = RGBColor(*_hex_to_rgb(accent))

    # Espace médian
    for _ in range(6):
        doc.add_paragraph()

    # Auteur
    if settings.get("include_author", True):
        author = settings.get("author", "")
        if author:
            author_para = doc.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_run = author_para.add_run(author)
            author_run.font.name = font_cfg["body"]
            author_run.font.size = Pt(13)
            author_run.bold = True
            author_run.font.color.rgb = RGBColor(*_hex_to_rgb(primary))

    # Organisation
    if settings.get("include_organization", True):
        org = settings.get("organization", "")
        if org:
            org_para = doc.add_paragraph()
            org_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            org_run = org_para.add_run(org)
            org_run.font.name = font_cfg["body"]
            org_run.font.size = Pt(11)
            org_run.font.color.rgb = RGBColor(*_hex_to_rgb(accent))

    # Date
    if settings.get("include_date", True):
        date_display = settings.get("date") or settings.get("_generated_date", "")
        if date_display:
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_run = date_para.add_run(date_display)
            date_run.font.name = font_cfg["body"]
            date_run.font.size = Pt(10)
            date_run.font.color.rgb = RGBColor(*_hex_to_rgb(gray))

    # Version
    version = settings.get("version", "")
    if version:
        ver_para = doc.add_paragraph()
        ver_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ver_run = ver_para.add_run(f"Version {version}")
        ver_run.font.name = font_cfg["body"]
        ver_run.font.size = Pt(9)
        ver_run.font.color.rgb = RGBColor(*_hex_to_rgb(gray))


def _add_toc_docx(
    doc: Document,
    headings: list[dict],
    settings: dict,
    font_cfg: dict,
) -> None:
    """Génère une table des matières statique à partir des titres extraits."""
    colors_theme = settings.get("theme_colors", {})
    primary = colors_theme.get("primary", "#1a1a1a")
    accent = colors_theme.get("accent", "#666666")

    # Titre de la TOC
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tr = toc_title.add_run("Table des matières")
    tr.font.name = font_cfg["heading"]
    tr.font.size = Pt(font_cfg["h2_size"])
    tr.bold = True
    tr.font.color.rgb = RGBColor(*_hex_to_rgb(primary))

    doc.add_paragraph()

    # Entrées par niveau
    for h in headings:
        level = h["level"]
        text = h["title"]

        if level == 3:
            indent = Inches(0.4)
            size = Pt(10)
        elif level == 2:
            indent = Inches(0.2)
            size = Pt(11)
        else:  # level 1
            indent = Inches(0.0)
            size = Pt(12)

        entry = doc.add_paragraph()
        entry.paragraph_format.left_indent = indent
        run = entry.add_run(text)
        run.font.name = font_cfg["body"]
        run.font.size = size
        run.font.color.rgb = RGBColor(*_hex_to_rgb(accent))


def _convert_markdown_to_docx_pub(
    md_text: str,
    doc: Document,
    font_cfg: dict,
    theme_colors: dict,
) -> None:
    """
    Variante de _convert_markdown_to_docx avec styles thématiques.
    Supporte H1/H2/H3, listes, citations, code, séparateurs, paragraphes.
    """
    primary = theme_colors.get("primary", "#1a1a1a")
    accent = theme_colors.get("accent", "#666666")

    lines = md_text.splitlines()
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Bloc de code
        if stripped.startswith("```"):
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            _add_code_block(doc, "\n".join(code_lines))
            i += 1
            continue

        # Séparateur
        if re.match(r"^-{3,}$", stripped) or re.match(r"^\*{3,}$", stripped):
            sep = doc.add_paragraph("─" * 50)
            sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # H1
        if stripped.startswith("# "):
            text = stripped[2:].strip()
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.name = font_cfg["heading"]
            run.font.size = Pt(font_cfg["h1_size"])
            run.bold = font_cfg["h1_bold"]
            run.font.color.rgb = RGBColor(*_hex_to_rgb(primary))
            i += 1
            continue

        # H2
        if stripped.startswith("## "):
            text = stripped[3:].strip()
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.name = font_cfg["heading"]
            run.font.size = Pt(font_cfg["h2_size"])
            run.bold = font_cfg["h2_bold"]
            run.font.color.rgb = RGBColor(*_hex_to_rgb(primary))
            i += 1
            continue

        # H3
        if stripped.startswith("### "):
            text = stripped[4:].strip()
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.name = font_cfg["heading"]
            run.font.size = Pt(font_cfg["h3_size"])
            run.bold = font_cfg["h3_bold"]
            run.font.color.rgb = RGBColor(*_hex_to_rgb(accent))
            i += 1
            continue

        # Liste à puces
        if re.match(r"^[*\-] ", stripped):
            text = re.sub(r"^[*\-] ", "", stripped)
            para = doc.add_paragraph(style="List Bullet")
            run = para.add_run(text)
            run.font.name = font_cfg["body"]
            run.font.size = Pt(font_cfg["body_size"])
            i += 1
            continue

        # Citation
        if stripped.startswith("> "):
            text = stripped[2:].strip()
            try:
                para = doc.add_paragraph(style="Quote")
            except KeyError:
                para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.name = font_cfg["body"]
            run.font.size = Pt(font_cfg["body_size"])
            run.italic = True
            i += 1
            continue

        # Ligne vide
        if stripped == "":
            i += 1
            continue

        # Paragraphe normal
        para = doc.add_paragraph()
        run = para.add_run(stripped)
        run.font.name = font_cfg["body"]
        run.font.size = Pt(font_cfg["body_size"])
        i += 1


def _settings_signature(settings: dict) -> str:
    """Signature des paramètres de gabarit (hors données dérivées)."""
    relevant = {
        k: v
        for k, v in settings.items()
        if not k.startswith("_")
        and k not in ("analysis", "theme_colors", "theme")
        and isinstance(v, (str, bool, int, float))
    }
    return hashlib.md5(
        json.dumps(relevant, sort_keys=True).encode()
    ).hexdigest()


def _add_image_cover_docx(
    doc: Document, cover_path: Path, settings: dict
) -> None:
    """Insère une image comme page de couverture dans le document Word."""
    page_size_key = settings.get("page_size", "letter")
    page_width, _ = _PAGE_SIZES_DOCX.get(page_size_key, _PAGE_SIZES_DOCX["letter"])
    page_margin = Inches(_PAGE_MARGINS_DOCX.get(page_size_key, 1.0))

    img_width = page_width - 2 * page_margin

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run()
    run.add_picture(str(cover_path), width=img_width)


def export_publication_docx(project_name: str) -> Path | None:
    """
    Génère final/document_publication.docx.

    Applique :
    - Page de couverture typographique
    - Table des matières statique
    - Styles selon template/theme/font_style
    - Marges selon page_size
    - En-têtes, pieds de page et numéros de page
    - Métadonnées DOCX complètes (titre, auteur, sujet, mots-clés, catégorie)

    Ne reconstruit pas si les sources n'ont pas changé.
    """
    from app.publication_template_service import (
        resolve_publication_settings,
        extract_headings,
    )

    final_dir = SORTIE_DIR / project_name / "final"
    pub_md_path = final_dir / "document_publication.md"
    final_md_path = final_dir / "document_final.md"
    pub_docx_path = final_dir / "document_publication.docx"

    if not pub_md_path.exists():
        print(
            f"[docx_export] document_publication.md introuvable "
            f"pour le projet : {project_name}"
        )
        return None

    final_content = (
        final_md_path.read_text(encoding="utf-8")
        if final_md_path.exists()
        else ""
    )
    settings = resolve_publication_settings(project_name, final_content)

    from app.cover_generation_service import get_cover_path

    cover_path = get_cover_path(project_name)
    cover_sig = _file_hash(cover_path) if cover_path else "none"

    pub_md_sig = _file_hash(pub_md_path)
    combined_sig = f"{pub_md_sig}:{_settings_signature(settings)}:{cover_sig}"

    state = load_project_state(project_name)
    pub_docx_state = state.get("publication", {}).get("docx", {})

    if (
        pub_docx_path.exists()
        and pub_docx_state.get("generated")
        and pub_docx_state.get("source_signature") == combined_sig
    ):
        print(f"[docx_export] DOCX publication déjà à jour : {pub_docx_path}")
        return pub_docx_path

    font_cfg = _FONT_CONFIGS_DOCX.get(
        settings["font_style"],
        _FONT_CONFIGS_DOCX["modern"],
    )
    theme_colors = settings["theme_colors"]

    doc = Document()

    # Métadonnées DOCX enrichies
    core = doc.core_properties
    core.title = settings.get("title", project_name)
    core.author = settings.get("author", "TranscriptionAI")
    core.subject = settings.get("subtitle", "Document de publication")
    core.keywords = settings.get("keywords", "")
    core.comments = settings.get("description", "")
    core.category = settings.get("category", "")

    # Format de page
    _apply_page_size_docx(doc, settings["page_size"])

    # En-têtes / pieds de page
    _add_headers_footers_docx(doc, settings, font_cfg)

    # Couverture
    if settings.get("include_cover", True):
        if cover_path:
            _add_image_cover_docx(doc, cover_path, settings)
        else:
            _add_cover_page_docx(doc, settings, font_cfg)
        doc.add_page_break()

    # Table des matières
    if settings.get("include_toc", True) and final_content:
        headings = extract_headings(final_content)
        if headings:
            _add_toc_docx(doc, headings, settings, font_cfg)
            doc.add_page_break()

    # Contenu principal (document_final.md)
    _convert_markdown_to_docx_pub(final_content, doc, font_cfg, theme_colors)

    pub_docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(pub_docx_path))

    updated_at = datetime.now().isoformat(timespec="seconds")

    if "publication" not in state:
        state["publication"] = {}

    state["publication"]["docx"] = {
        "generated": True,
        "path": str(pub_docx_path),
        "source_signature": combined_sig,
        "updated_at": updated_at,
    }

    save_project_state(project_name, state)

    print(f"[docx_export] DOCX publication généré : {pub_docx_path}")
    return pub_docx_path
